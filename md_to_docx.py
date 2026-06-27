#!/usr/bin/env python3
"""
Simple Markdown to DOCX converter tailored for the workshop analysis file.
Creates a clean, professional .docx with proper headings, the main table, bullet lists, and inline bold.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell (e.g. 'D9E2F3' for light blue)."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Set reasonable page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Read the markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip empty lines but add a small paragraph break sometimes
        if not line.strip():
            # Only add spacing paragraph if previous was not a heading or list
            if i > 0 and lines[i-1].strip() and not lines[i-1].strip().startswith(('#', '-', '*', '|')):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('─' * 60)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Markdown table detection - only the main analysis table (more strict to avoid false positives from raw data)
        if '|' in line and line.count('|') >= 3 and 'Dokument' in line and 'Ord 1' in line:
            # Start collecting table
            table_lines = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].rstrip('\n')
                if '|' in next_line and next_line.count('|') >= 3:
                    table_lines.append(next_line)
                    j += 1
                else:
                    break

            # Parse table - skip separator rows (those with lots of ---)
            rows_data = []
            for tl in table_lines:
                if '---' in tl and tl.count('-') > 5:
                    continue  # skip markdown table separator row
                # Split and clean cells, skip leading/trailing empty from |
                cells = [c.strip() for c in tl.split('|')]
                cells = [c for c in cells if c]  # remove empties from start/end
                if cells:
                    rows_data.append(cells)

            if rows_data:
                num_rows = len(rows_data)
                num_cols = len(rows_data[0])
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'

                # Header row styling (light blue-ish)
                for c_idx, cell_text in enumerate(rows_data[0]):
                    cell = table.rows[0].cells[c_idx]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                    set_cell_shading(cell, 'D6EAF8')  # light blue header

                # Data rows
                for r_idx in range(1, num_rows):
                    for c_idx, cell_text in enumerate(rows_data[r_idx]):
                        if c_idx < num_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

                # Add a bit of space after table
                doc.add_paragraph()

            i = j
            continue

        # Bullet lists (- or *)
        if line.strip().startswith(('- ', '* ')):
            bullet_text = line.strip()[2:].strip()
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
            continue

        # Numbered lists (simple support)
        if re.match(r'^\d+\.\s', line.strip()):
            num_text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(num_text, style='List Number')
            i += 1
            continue

        # Regular paragraph with possible **bold** inline
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

        # Split on **...** for bold
        parts = re.split(r'(\*\*[^*]+\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

        i += 1

    # Save
    doc.save(docx_path)
    print(f"Successfully created: {docx_path}")
    return docx_path

if __name__ == "__main__":
    md_file = "workshop_8_dokument_analys_sammanslagen.md"
    docx_file = "workshop_8_dokument_analys_sammanslagen.docx"

    base_dir = os.path.dirname(os.path.abspath(__file__)) or "."
    md_path = os.path.join(base_dir, md_file)
    docx_path = os.path.join(base_dir, docx_file)

    if not os.path.exists(md_path):
        print(f"ERROR: Markdown file not found: {md_path}")
    else:
        convert_md_to_docx(md_path, docx_path)